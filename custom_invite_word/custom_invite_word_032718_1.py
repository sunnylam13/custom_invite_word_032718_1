# -*- coding: utf-8 -*-

#! python3

# USAGE
# python3 custom_invite_word_032718_1.py "GUESTNAMES.TXT"
# python3 custom_invite_word_032718_1.py "../tests/guests.txt"

import docx,sys,re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import logging
logging.basicConfig(level=logging.DEBUG, format=" %(asctime)s - %(levelname)s - %(message)s")
# logging.disable(logging.CRITICAL)

#####################################
# GUEST INFORMATION
#####################################

# gather guest information

guest_txt = sys.argv[1]

logging.debug( 'Guest list text file input is:  %s' % (guest_txt) )

# remove carriage return
# https://regexr.com/3n1el
line_break_regex1 = re.compile("(\\n)" + "$")

def remove_carriage_return(list):
	new_list = []
	for item in list:
		new_item = line_break_regex1.sub('',item)
		new_list.append(new_item)
	return new_list

def guest_list_maker(guest_txt):
	guest_list = []
	# open the guest list text file
	fileObj = open(guest_txt)
	guest_list = fileObj.readlines() # this should return a list
	guest_list = remove_carriage_return(guest_list)
	logging.debug( 'The guest list extracted from text is:  ' )
	logging.debug( guest_list )
	# return the guest_list to outside the function for use
	return guest_list

#####################################
# END GUEST INFORMATION
#####################################

#####################################
# WORD DOCUMENT
#####################################

# create new Word document

doc = docx.Document()

# create styles

## character styling

styles = doc.styles # whole doc's styles object

style_italic_emph_1 = styles.add_style('GuestItalics', WD_STYLE_TYPE.CHARACTER) # use 'GuestItalics' as the name when applying below!
logging.debug( "The style name created is:  %s" % (style_italic_emph_1.name) )
logging.debug( "The style type created is:  %s" % (style_italic_emph_1.name) )
logging.debug( style_italic_emph_1.type )
font_style_italic_emph_1 = style_italic_emph_1.font
font_style_italic_emph_1.name = 'Calibri'
font_style_italic_emph_1.size = Pt(14)
font_style_italic_emph_1.italic = True

style_italic_underline_1 = styles.add_style('ItalUnder1', WD_STYLE_TYPE.CHARACTER) # use 'ItalUnder1' as the name when applying below!  also registers it in Word doc styles list
logging.debug( "The style name created is:  %s" % (style_italic_underline_1.name) )
logging.debug( "The style type created is:  %s" % (style_italic_underline_1.name) )
style_italic_underline_1 = style_italic_underline_1.font
style_italic_underline_1.name = 'Times New Roman'
style_italic_underline_1.size = Pt(11)
style_italic_underline_1.italic = True
style_italic_underline_1.underline = True

## paragraph styling

center_style_1 = doc.styles.add_style('Center1', WD_STYLE_TYPE.PARAGRAPH)  # use 'Center1' as the name when applying below!  also registers it in Word doc styles list
center_para_1 = center_style_1.paragraph_format
center_para_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
# center_para_1.base_style = styles['Subtitle']
center_para_1 = center_style_1.font
center_para_1.name = 'Calibri'
center_para_1.size = Pt(14)
center_para_1.italic = True

def construct_invite_1(doc,guest_list):

	counter = 0

	for name in guest_list:
		doc.add_paragraph('It would be a pleasure to have the company of')
		logging.debug(doc.paragraphs[0].text)
		doc.paragraphs[0 + counter].style = 'Subtitle'
		# doc.paragraphs[0 + counter].style = doc.styles['Center1']
		doc.paragraphs[0 + counter].alignment = WD_ALIGN_PARAGRAPH.CENTER
		logging.debug('Styles applied')
		
		# insert the guest's name
		doc.add_paragraph(name)
		logging.debug(doc.paragraphs[1 + counter].text)
		doc.paragraphs[1 + counter].style = 'Heading 1'
		doc.paragraphs[1 + counter].alignment = WD_ALIGN_PARAGRAPH.CENTER
		logging.debug('Styles applied')
		
		doc.add_paragraph().add_run('at')
		doc.paragraphs[2 + counter].style = doc.styles['Center1']
		doc.paragraphs[2 + counter].add_run(' 11010 Memory Lane on the Evening of') # add to the same paragraph
		logging.debug(doc.paragraphs[2 + counter].runs[0].text)
		logging.debug(doc.paragraphs[2 + counter].runs[1].text)
		# doc.paragraphs[2 + counter].runs[0].italic = True # 'at'
		doc.paragraphs[2 + counter].runs[0].style = doc.styles['ItalUnder1'] # 'at'
		doc.paragraphs[2 + counter].runs[1].style = doc.styles['GuestItalics'] # '11010 Memory Lane on the Evening of'
		logging.debug('Styles applied')

		doc.add_paragraph('April 1st')
		logging.debug(doc.paragraphs[3 + counter].text)
		logging.debug(doc.paragraphs[3 + counter].runs[0].text)
		doc.paragraphs[3 + counter].style = 'Normal'
		doc.paragraphs[3 + counter].alignment = WD_ALIGN_PARAGRAPH.CENTER
		logging.debug('Styles applied')

		# doc.add_paragraph("at 7 o'clock")
		doc.add_paragraph().add_run('at')
		doc.paragraphs[4 + counter].style = doc.styles['Center1']
		doc.paragraphs[4 + counter].add_run(" 7 o'clock") # add to the same paragraph
		logging.debug(doc.paragraphs[4 + counter].runs[0].text)
		logging.debug(doc.paragraphs[4 + counter].runs[1].text)
		# doc.paragraphs[4 + counter].runs[0].underline.italic = True # 'at'
		doc.paragraphs[4 + counter].runs[0].style = doc.styles['ItalUnder1'] # 'at'
		doc.paragraphs[4 + counter].runs[1].style = doc.styles['GuestItalics'] # '7 o'clock'
		logging.debug('Styles applied')

		doc.add_page_break() # add page break

		counter += 6 # increment by 6 positions so that you can style the next guest properly, apparently we need to count the line break so it's not 5
		
#####################################
# END WORD DOCUMENT
#####################################

#####################################
# EXECUTION
#####################################

construct_invite_1( doc,guest_list_maker(guest_txt) )

# save the final word doc

doc.save('custom_invites_f.docx')

#####################################
# END EXECUTION
#####################################

